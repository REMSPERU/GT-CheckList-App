import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, text, url):
    """
    Adds a hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    # Color (Blue)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0056B3')
    rPr.append(color)
    
    new_run.append(rPr)
    
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def parse_inline_formatting(paragraph, text):
    """
    Parses a string of text containing **bold**, `code`, and [text](url) links,
    and adds them as properly formatted runs to the paragraph.
    """
    # Pattern to match:
    # 1. Markdown Links: [text](url) -> Group 1 & 2
    # 2. Bold text: **text** -> Group 3
    # 3. Inline code: `text` -> Group 4
    # 4. Plain text: any other characters -> Group 5
    pattern = r'\[([^\]]+)\]\(([^\)]+)\)|(\*\*.*?\*\*)|(`.*?`)|([^\*`\[]+|[\*`\[](?![^\*`\[]*[\*`\]]))'
    matches = re.finditer(pattern, text)
    
    for match in matches:
        link_text = match.group(1)
        link_url = match.group(2)
        bold_text = match.group(3)
        code_text = match.group(4)
        plain_text = match.group(5)
        
        if link_text and link_url:
            add_hyperlink(paragraph, link_text, link_url)
        elif bold_text:
            content = bold_text[2:-2] # Strip **
            run = paragraph.add_run(content)
            run.bold = True
        elif code_text:
            content = code_text[1:-1] # Strip `
            run = paragraph.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            # Give inline code a dark gray color
            run.font.color.rgb = RGBColor(180, 40, 40)
        elif plain_text:
            paragraph.add_run(plain_text)

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    root_dir = os.path.abspath(os.path.join(script_dir, '..'))
    
    md_file_path = os.path.join(root_dir, 'docs', 'informe_funcionalidades.md')
    docx_file_path = os.path.join(root_dir, 'docs', 'informe_funcionalidades.docx')
    
    print(f"Reading Markdown file from: {md_file_path}")
    if not os.path.exists(md_file_path):
        print(f"Error: {md_file_path} does not exist.")
        return
        
    # Read Markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    doc = docx.Document()
    
    # Configure page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Configure base style (Normal paragraph font)
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(33, 37, 41) # Dark gray body text
    
    in_code_block = False
    code_block_text = []
    
    print("Converting Markdown to Docx...")
    for line_num, line in enumerate(lines, 1):
        stripped_line = line.strip()
        
        # Code block toggle
        if stripped_line.startswith('```'):
            if in_code_block:
                # End code block
                in_code_block = False
                code_text = '\n'.join(code_block_text)
                
                # Add code block with gray background style
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(50, 50, 50)
                
                code_block_text = []
            else:
                # Start code block
                in_code_block = True
            continue
            
        if in_code_block:
            code_block_text.append(line.rstrip('\n'))
            continue
            
        # Headers
        if stripped_line.startswith('# '):
            title_text = stripped_line[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(10, 30, 80) # Dark blue title
            
        elif stripped_line.startswith('## '):
            header_text = stripped_line[3:]
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            
            run = p.add_run(header_text)
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(20, 50, 120)
            
        elif stripped_line.startswith('### '):
            header_text = stripped_line[4:]
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
            run = p.add_run(header_text)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(30, 70, 150)
            
        elif stripped_line.startswith('#### '):
            header_text = stripped_line[5:]
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            
            run = p.add_run(header_text)
            run.bold = True
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(50, 50, 50)
            
        # Horizontal rule
        elif stripped_line == '---':
            # Add a visual divider (just a simple empty paragraph with space)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            
        # Bullet list item
        elif stripped_line.startswith('* ') or stripped_line.startswith('- '):
            bullet_text = stripped_line[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            parse_inline_formatting(p, bullet_text)
            
        # Numbered list item
        elif re.match(r'^\d+\.\s', stripped_line):
            match = re.match(r'^\d+\.\s(.*)', stripped_line)
            num_text = match.group(1)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            parse_inline_formatting(p, num_text)
            
        # Normal paragraph (excluding empty lines)
        elif stripped_line:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            parse_inline_formatting(p, stripped_line)
            
    # Save the docx file
    try:
        doc.save(docx_file_path)
        print(f"Word document generated successfully at: {docx_file_path}")
    except PermissionError:
        print("\n" + "="*80)
        print("ERROR DE PERMISOS: No se pudo guardar el archivo Word (.docx).")
        print("Esto ocurre generalmente porque el archivo está abierto en Microsoft Word u otro visor.")
        print("Por favor, CIERRA el documento en Microsoft Word y vuelve a ejecutar el script.")
        print("="*80 + "\n")
        import sys
        sys.exit(1)

if __name__ == "__main__":
    main()
